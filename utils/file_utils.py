import logging
from PIL import Image
import PyPDF2
import docx
import base64
import io
from pdf2image import convert_from_bytes
from charset_normalizer import detect

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('file_utils.log'),
        logging.StreamHandler()
    ]
)

class FileProcessingError(Exception):
    """Custom exception for file processing errors."""
    pass

def extract_text_from_pdf(file, max_size_mb=10):
    """
    Extract text from a PDF file.

    Args:
        file: File object to process.
        max_size_mb (int): Maximum allowed file size in MB.

    Returns:
        str: Extracted text from the PDF.

    Raises:
        FileProcessingError: If the file is too large, corrupted, or cannot be processed.
    """
    try:
        # Check file size
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        reader = PyPDF2.PdfReader(file)
        if reader.is_encrypted:
            raise FileProcessingError("PDF is password-protected.")
        
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted
        if not text.strip():
            logging.warning("No text extracted from PDF.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing PDF: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing PDF: {str(e)}")
        raise FileProcessingError(f"Failed to process PDF: {str(e)}")

def extract_text_from_docx(file, max_size_mb=10):
    """
    Extract text from a DOCX file.

    Args:
        file: File object to process.
        max_size_mb (int): Maximum allowed file size in MB.

    Returns:
        str: Extracted text from the DOCX.

    Raises:
        FileProcessingError: If the file is too large, corrupted, or cannot be processed.
    """
    try:
        # Check file size
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        if not text.strip():
            logging.warning("No text extracted from DOCX.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing DOCX: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing DOCX: {str(e)}")
        raise FileProcessingError(f"Failed to process DOCX: {str(e)}")

def extract_text_from_txt(file, max_size_mb=10):
    """
    Extract text from a TXT file.

    Args:
        file: File object to process.
        max_size_mb (int): Maximum allowed file size in MB.

    Returns:
        str: Extracted text from the TXT file.

    Raises:
        FileProcessingError: If the file is too large, corrupted, or cannot be processed.
    """
    try:
        # Check file size
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        content = file.read()
        detected = detect(content)
        encoding = detected.get("encoding", "utf-8")
        text = content.decode(encoding)
        if not text.strip():
            logging.warning("No text extracted from TXT.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing TXT: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing TXT: {str(e)}")
        raise FileProcessingError(f"Failed to process TXT: {str(e)}")

def pdf_to_image(file, max_size_mb=10):
    """
    Convert a PDF file to an image.

    Args:
        file: File object to process.
        max_size_mb (int): Maximum allowed file size in MB.

    Returns:
        BytesIO: Buffer containing the first page of the PDF as a PNG image, or None if conversion fails.

    Raises:
        FileProcessingError: If the file is too large, corrupted, or cannot be processed.
    """
    try:
        # Check file size
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        images = convert_from_bytes(file.read())
        if images:
            buffered = io.BytesIO()
            images[0].save(buffered, format="PNG")
            buffered.seek(0)
            return buffered
        else:
            raise FileProcessingError("No pages found in PDF.")
    except FileProcessingError as e:
        logging.error(f"Error converting PDF to image: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error converting PDF to image: {str(e)}")
        raise FileProcessingError(f"Failed to convert PDF to image: {str(e)}")

def get_image_parts(uploaded_file, max_size_mb=10):
    """
    Convert an image file to a format suitable for Gemini API.

    Args:
        uploaded_file: File object to process.
        max_size_mb (int): Maximum allowed file size in MB.

    Returns:
        dict: Dictionary containing the image data in a format for the Gemini API.

    Raises:
        FileProcessingError: If the file is too large, corrupted, or cannot be processed.
    """
    try:
        # Check file size
        uploaded_file.seek(0, io.SEEK_END)
        file_size = uploaded_file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        uploaded_file.seek(0)

        image = Image.open(uploaded_file)
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        image_bytes = buffered.getvalue()
        return {
            "inline_data": {
                "mime_type": "image/png",
                "data": base64.b64encode(image_bytes).decode("utf-8"),
            }
        }
    except FileProcessingError as e:
        logging.error(f"Error processing image: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing image: {str(e)}")
        raise FileProcessingError(f"Failed to process image: {str(e)}")
    
    import logging
from PIL import Image
import PyPDF2
import docx
import base64
import io
from pdf2image import convert_from_bytes
from charset_normalizer import detect

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('file_utils.log'),
        logging.StreamHandler()
    ]
)

class FileProcessingError(Exception):
    """Custom exception for file processing errors."""
    pass

def extract_text_from_pdf(file, max_size_mb=10):
    try:
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        reader = PyPDF2.PdfReader(file)
        if reader.is_encrypted:
            raise FileProcessingError("PDF is password-protected.")
        
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted
        if not text.strip():
            logging.warning("No text extracted from PDF.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing PDF: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing PDF: {str(e)}")
        raise FileProcessingError(f"Failed to process PDF: {str(e)}")

def extract_text_from_docx(file, max_size_mb=10):
    try:
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        if not text.strip():
            logging.warning("No text extracted from DOCX.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing DOCX: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing DOCX: {str(e)}")
        raise FileProcessingError(f"Failed to process DOCX: {str(e)}")

def extract_text_from_txt(file, max_size_mb=10):
    try:
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        content = file.read()
        detected = detect(content)
        encoding = detected.get("encoding", "utf-8")
        text = content.decode(encoding)
        if not text.strip():
            logging.warning("No text extracted from TXT.")
        return text
    except FileProcessingError as e:
        logging.error(f"Error processing TXT: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing TXT: {str(e)}")
        raise FileProcessingError(f"Failed to process TXT: {str(e)}")

def pdf_to_image(file, max_size_mb=10):
    try:
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        file.seek(0)

        images = convert_from_bytes(file.read())
        if images:
            buffered = io.BytesIO()
            images[0].save(buffered, format="PNG")
            buffered.seek(0)
            return buffered
        else:
            raise FileProcessingError("No pages found in PDF.")
    except FileProcessingError as e:
        logging.error(f"Error converting PDF to image: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error converting PDF to image: {str(e)}")
        raise FileProcessingError(f"Failed to convert PDF to image: {str(e)}")

def get_image_parts(uploaded_file, max_size_mb=10):
    try:
        uploaded_file.seek(0, io.SEEK_END)
        file_size = uploaded_file.tell()
        if file_size > max_size_mb * 1024 * 1024:
            raise FileProcessingError(f"File size exceeds {max_size_mb}MB limit.")
        uploaded_file.seek(0)

        image = Image.open(uploaded_file)
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        image_bytes = buffered.getvalue()
        return {
            "inline_data": {
                "mime_type": "image/png",
                "data": base64.b64encode(image_bytes).decode("utf-8"),
            }
        }
    except FileProcessingError as e:
        logging.error(f"Error processing image: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Unexpected error processing image: {str(e)}")
        raise FileProcessingError(f"Failed to process image: {str(e)}")